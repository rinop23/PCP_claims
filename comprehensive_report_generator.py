"""
Comprehensive Report Generator for Milberg Monthly Reports
Generates detailed analysis with tables, graphs, and insights from Word documents
"""

import os
from typing import Dict, Any, List
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import io
from PIL import Image


class ComprehensiveReportGenerator:
    """Generate comprehensive analysis reports from extracted data"""

    def __init__(self):
        self.charts_dir = "temp_charts"
        os.makedirs(self.charts_dir, exist_ok=True)

    def generate_full_report(self, extracted_data: Dict[str, Any], output_path: str) -> str:
        """
        Generate a comprehensive Word document report with all analysis
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        self._add_title(doc, "Comprehensive Claims Analysis Report")
        self._add_subtitle(doc, f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        doc.add_paragraph()

        # Executive Summary
        self._add_executive_summary(doc, extracted_data)

        # Portfolio Overview
        self._add_portfolio_overview(doc, extracted_data)

        # Claims Analysis by Lender
        self._add_lender_analysis(doc, extracted_data)

        # Detailed Claims Table
        self._add_claims_table(doc, extracted_data)

        # Charts and Visualizations
        self._add_visualizations(doc, extracted_data)

        # FCA Compliance Summary (if available)
        if 'claim_eligibility' in extracted_data:
            self._add_fca_compliance_section(doc, extracted_data)

        # Financial Analysis
        self._add_financial_analysis(doc, extracted_data)

        # Recommendations
        self._add_recommendations(doc, extracted_data)

        # Save document
        doc.save(output_path)
        print(f"[Success] Comprehensive report saved to: {output_path}")

        # Clean up temporary charts
        self._cleanup_charts()

        return output_path

    def _add_title(self, doc, text: str):
        """Add formatted title"""
        heading = doc.add_heading(text, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(31, 78, 121)

    def _add_subtitle(self, doc, text: str):
        """Add formatted subtitle"""
        para = doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(89, 89, 89)

    def _add_executive_summary(self, doc, data: Dict[str, Any]):
        """Add executive summary section"""
        doc.add_heading("Executive Summary", level=1)

        portfolio = data.get('portfolio_summary', {})
        total_claims = data.get('total_claims', 0)
        total_claimants = portfolio.get('total_claimants', 0)

        summary_text = f"""
This report provides a comprehensive analysis of the current claims portfolio.
Key highlights include:

• Total Claims in Portfolio: {total_claims:,}
• Total Unique Claimants: {total_claimants:,}
• Claims Submitted to Scheme: {portfolio.get('total_claims_submitted', 0):,}
• Extraction Method: {data.get('extraction_method', 'Unknown')}
• Report Source: {os.path.basename(data.get('source_file', 'Unknown'))}
        """

        doc.add_paragraph(summary_text.strip())
        doc.add_paragraph()

    def _add_portfolio_overview(self, doc, data: Dict[str, Any]):
        """Add portfolio overview with key metrics"""
        doc.add_heading("Portfolio Overview", level=1)

        portfolio = data.get('portfolio_summary', {})

        # Create metrics table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'

        # Add metrics
        metrics = [
            ('Total Claims', f"{data.get('total_claims', 0):,}"),
            ('Total Claimants', f"{portfolio.get('total_claimants', 0):,}"),
            ('Claims Submitted', f"{portfolio.get('total_claims_submitted', 0):,}"),
            ('Total Claim Value', f"£{portfolio.get('total_claim_value', 0):,.2f}"),
            ('Total Funded', f"£{portfolio.get('total_funded', 0):,.2f}"),
            ('Report Type', portfolio.get('report_type', 'N/A')),
        ]

        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)

        doc.add_paragraph()

    def _add_lender_analysis(self, doc, data: Dict[str, Any]):
        """Add analysis by lender/defendant"""
        doc.add_heading("Claims Distribution by Lender", level=1)

        claims = data.get('claims', [])

        # Count claims by defendant
        lender_counts = {}
        for claim in claims:
            defendant = claim.get('defendant', 'Unknown')
            lender_counts[defendant] = lender_counts.get(defendant, 0) + 1

        # Sort by count
        sorted_lenders = sorted(lender_counts.items(), key=lambda x: x[1], reverse=True)

        # Create table
        if sorted_lenders:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Lender / Defendant'
            hdr_cells[1].text = 'Number of Claims'
            hdr_cells[2].text = '% of Total'

            total = sum(lender_counts.values())

            # Add top 20 lenders
            for lender, count in sorted_lenders[:20]:
                row_cells = table.add_row().cells
                row_cells[0].text = lender
                row_cells[1].text = str(count)
                row_cells[2].text = f"{(count/total*100):.1f}%"

            if len(sorted_lenders) > 20:
                doc.add_paragraph(f"\n(Showing top 20 of {len(sorted_lenders)} total lenders)")

        doc.add_paragraph()

    def _add_claims_table(self, doc, data: Dict[str, Any]):
        """Add detailed claims table"""
        doc.add_heading("Detailed Claims Information", level=1)

        claims = data.get('claims', [])

        if claims:
            # Show first 50 claims in detail
            sample_claims = claims[:50]

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Claim ID'
            hdr_cells[1].text = 'Defendant'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Claim Amount'
            hdr_cells[4].text = 'Funded Amount'

            # Add claims
            for claim in sample_claims:
                row_cells = table.add_row().cells
                row_cells[0].text = claim.get('claim_id', 'N/A')[:30]
                row_cells[1].text = claim.get('defendant', 'N/A')[:40]
                row_cells[2].text = claim.get('status', 'N/A')
                claim_amt = claim.get('claim_amount', 0)
                funded_amt = claim.get('funded_amount', 0)
                row_cells[3].text = f"£{claim_amt:,.2f}" if claim_amt > 0 else "TBD"
                row_cells[4].text = f"£{funded_amt:,.2f}" if funded_amt > 0 else "TBD"

            if len(claims) > 50:
                doc.add_paragraph(f"\n(Showing 50 of {len(claims):,} total claims)")

        doc.add_paragraph()

    def _add_visualizations(self, doc, data: Dict[str, Any]):
        """Add charts and visualizations"""
        doc.add_heading("Visual Analysis", level=1)

        claims = data.get('claims', [])

        # Chart 1: Top 10 Lenders by Claims Count
        self._add_top_lenders_chart(doc, claims)

        # Chart 2: Claims Status Distribution
        self._add_status_distribution_chart(doc, claims)

        doc.add_paragraph()

    def _add_top_lenders_chart(self, doc, claims: List[Dict]):
        """Create and add top lenders bar chart"""
        # Count claims by lender
        lender_counts = {}
        for claim in claims:
            defendant = claim.get('defendant', 'Unknown')
            lender_counts[defendant] = lender_counts.get(defendant, 0) + 1

        # Get top 10
        sorted_lenders = sorted(lender_counts.items(), key=lambda x: x[1], reverse=True)[:10]

        if sorted_lenders:
            lenders = [l[0][:30] for l in sorted_lenders]  # Truncate long names
            counts = [l[1] for l in sorted_lenders]

            # Create chart
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.barh(lenders, counts, color='#1f4e79')
            ax.set_xlabel('Number of Claims', fontsize=12)
            ax.set_title('Top 10 Lenders by Claims Volume', fontsize=14, fontweight='bold')
            ax.invert_yaxis()

            # Add value labels
            for i, v in enumerate(counts):
                ax.text(v + 0.5, i, str(v), va='center')

            plt.tight_layout()

            # Save and add to document
            chart_path = os.path.join(self.charts_dir, 'top_lenders.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            doc.add_paragraph("Top 10 Lenders by Claims Volume:", style='Heading 2')
            doc.add_picture(chart_path, width=Inches(6))
            doc.add_paragraph()

    def _add_status_distribution_chart(self, doc, claims: List[Dict]):
        """Create and add status distribution pie chart"""
        # Count by status
        status_counts = {}
        for claim in claims:
            status = claim.get('status', 'Unknown')
            status_counts[status] = status_counts.get(status, 0) + 1

        if status_counts:
            labels = list(status_counts.keys())
            sizes = list(status_counts.values())
            colors = ['#1f4e79', '#4472c4', '#5b9bd5', '#70ad47', '#ffc000']

            # Create chart
            fig, ax = plt.subplots(figsize=(8, 8))
            wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%',
                                               colors=colors[:len(labels)], startangle=90)

            # Improve text readability
            for text in texts:
                text.set_fontsize(11)
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(10)

            ax.set_title('Claims Status Distribution', fontsize=14, fontweight='bold')

            # Save and add to document
            chart_path = os.path.join(self.charts_dir, 'status_distribution.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            doc.add_paragraph("Claims Status Distribution:", style='Heading 2')
            doc.add_picture(chart_path, width=Inches(5))
            doc.add_paragraph()

    def _add_fca_compliance_section(self, doc, data: Dict[str, Any]):
        """Add FCA compliance analysis section"""
        doc.add_heading("FCA Compliance Analysis", level=1)

        claim_eligibility = data.get('claim_eligibility', {})

        if claim_eligibility:
            # Count eligible vs ineligible
            eligible_count = 0
            ineligible_count = 0
            needs_review_count = 0

            for claim_id, eligibility in claim_eligibility.items():
                if hasattr(eligibility, 'is_eligible'):
                    if eligibility.is_eligible:
                        eligible_count += 1
                    else:
                        ineligible_count += 1
                elif isinstance(eligibility, dict):
                    if eligibility.get('is_eligible'):
                        eligible_count += 1
                    else:
                        ineligible_count += 1

            total_checked = eligible_count + ineligible_count + needs_review_count

            doc.add_paragraph(f"Total Claims Checked: {total_checked}")
            doc.add_paragraph(f"Eligible Claims: {eligible_count} ({eligible_count/total_checked*100:.1f}%)" if total_checked > 0 else "Eligible Claims: 0")
            doc.add_paragraph(f"Ineligible Claims: {ineligible_count} ({ineligible_count/total_checked*100:.1f}%)" if total_checked > 0 else "Ineligible Claims: 0")

        doc.add_paragraph()

    def _add_financial_analysis(self, doc, data: Dict[str, Any]):
        """Add financial analysis section"""
        doc.add_heading("Financial Analysis", level=1)

        portfolio = data.get('portfolio_summary', {})
        claims = data.get('claims', [])

        total_claim_value = portfolio.get('total_claim_value', 0)
        total_funded = portfolio.get('total_funded', 0)

        # Calculate metrics
        avg_claim_value = total_claim_value / len(claims) if claims else 0
        avg_funded = total_funded / len(claims) if claims else 0

        analysis_text = f"""
Key Financial Metrics:

• Total Potential Claim Value: £{total_claim_value:,.2f}
• Total Funding Provided: £{total_funded:,.2f}
• Average Claim Value: £{avg_claim_value:,.2f}
• Average Funding per Claim: £{avg_funded:,.2f}
• Number of Claims: {len(claims):,}

Note: Some values may show £0.00 if detailed financial data is pending submission
or redress calculations from the FCA scheme.
        """

        doc.add_paragraph(analysis_text.strip())
        doc.add_paragraph()

    def _add_recommendations(self, doc, data: Dict[str, Any]):
        """Add recommendations section"""
        doc.add_heading("Recommendations & Next Steps", level=1)

        claims = data.get('claims', [])
        portfolio = data.get('portfolio_summary', {})

        recommendations = []

        # Check if claims are submitted
        submitted = portfolio.get('total_claims_submitted', 0)
        total = len(claims)

        if submitted < total:
            recommendations.append(
                f"• Prioritize submission of remaining {total - submitted:,} claims to the FCA redress scheme"
            )

        # Check for OpenAI usage
        if data.get('extraction_method') == 'fallback':
            recommendations.append(
                "• Consider setting up OpenAI API key for enhanced document analysis (see OPENAI_SETUP.md)"
            )

        # General recommendations
        recommendations.extend([
            "• Monitor claim status updates from respondents regularly",
            "• Track redress calculation timeline for accepted claims",
            "• Maintain detailed records of all scheme communications",
            "• Review monthly reports for portfolio performance trends"
        ])

        for rec in recommendations:
            doc.add_paragraph(rec)

        doc.add_paragraph()

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph("End of Report")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.runs[0]
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _cleanup_charts(self):
        """Clean up temporary chart files"""
        try:
            import shutil
            if os.path.exists(self.charts_dir):
                shutil.rmtree(self.charts_dir)
        except Exception as e:
            print(f"[Warning] Could not clean up charts directory: {e}")
