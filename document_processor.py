"""
Document Processing Module for PCP Claim Reports

Handles extraction of structured data from law firm reports in various formats
"""

import re
import sys
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
import openpyxl
from openpyxl import load_workbook
import pandas as pd
from docx import Document
import os

# Set stdout encoding to UTF-8 for Windows compatibility
if sys.platform == 'win32':
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    else:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')


class DocumentProcessor:
    """Process and extract data from legal documents"""

    def __init__(self):
        self.extraction_patterns = self._compile_patterns()

    def call_openai_llm(self, prompt: str, model: str = "gpt-4o", max_tokens: int = 4000, temperature: float = 0.2) -> str:
        """
        Call OpenAI ChatCompletion API with a prompt and return the response text.
        Compatible with openai>=1.0.0
        """
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"[OpenAI API Error] {e}")
            return f"[OpenAI API Error: {str(e)}]"
    
    def _compile_patterns(self) -> Dict[str, re.Pattern]:
        """Compile regex patterns for data extraction"""
        return {
            "claim_id": re.compile(r"(?:Claim|Case|Reference)\s*(?:ID|Number|Ref)?:?\s*([A-Z0-9\-/]+)", re.IGNORECASE),
            "claimant": re.compile(r"Claimant(?:\s+Name)?:?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)", re.IGNORECASE),
            "defendant": re.compile(r"Defendant:?\s*([A-Za-z\s&,]+(?:Ltd|Limited|PLC|plc)?)", re.IGNORECASE),
            "claim_amount": re.compile(r"(?:Claim|Total|Amount)\s*(?:Value|Amount)?:?\s*£?\s*([\d,]+(?:\.\d{2})?)", re.IGNORECASE),
            "funded_amount": re.compile(r"(?:Funded|Funding|Advanced)\s*(?:Amount)?:?\s*£?\s*([\d,]+(?:\.\d{2})?)", re.IGNORECASE),
            "date": re.compile(r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(\d{4}-\d{2}-\d{2})"),
            "status": re.compile(r"Status:?\s*(Pending|In Progress|Settled|Rejected|Litigation)", re.IGNORECASE),
            "email": re.compile(r"[\w\.-]+@[\w\.-]+\.\w+"),
            "phone": re.compile(r"(?:\+44|0)[\d\s-]{10,}"),
        }
    
    def extract_from_text(self, text: str) -> Dict[str, Any]:
        """
        Extract structured data from text content
        Handles law firm reports in text format
        """
        extracted_data = {}
        
        # Extract claim ID
        match = self.extraction_patterns["claim_id"].search(text)
        if match:
            extracted_data["claim_id"] = match.group(1).strip()
        
        # Extract claimant
        match = self.extraction_patterns["claimant"].search(text)
        if match:
            extracted_data["claimant_name"] = match.group(1).strip()
        
        # Extract defendant
        match = self.extraction_patterns["defendant"].search(text)
        if match:
            extracted_data["defendant"] = match.group(1).strip()
        
        # Extract claim amount
        match = self.extraction_patterns["claim_amount"].search(text)
        if match:
            amount_str = match.group(1).replace(",", "")
            extracted_data["claim_amount"] = float(amount_str)
        
        # Extract funded amount
        match = self.extraction_patterns["funded_amount"].search(text)
        if match:
            amount_str = match.group(1).replace(",", "")
            extracted_data["funded_amount"] = float(amount_str)
        
        # Extract status
        match = self.extraction_patterns["status"].search(text)
        if match:
            extracted_data["status"] = match.group(1).lower().replace(" ", "_")
        
        # Extract dates
        dates = self.extraction_patterns["date"].findall(text)
        if dates:
            # Take the first date found as claim date
            date_str = dates[0][0] if dates[0][0] else dates[0][1]
            extracted_data["claim_date"] = self._normalize_date(date_str)
        
        # Extract law firm name (usually in header or footer)
        law_firm_patterns = [
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Solicitors|LLP|Limited|Law)",
            r"(?:From|Law Firm):\s*([A-Za-z\s&]+(?:LLP|Limited|Solicitors))"
        ]
        for pattern in law_firm_patterns:
            match = re.search(pattern, text)
            if match:
                extracted_data["law_firm"] = match.group(1).strip()
                break
        
        # Identify documentation types mentioned
        doc_keywords = {
            "claim_form": ["claim form", "particulars of claim"],
            "lfa_signed": ["funding agreement", "LFA signed", "litigation funding"],
            "insurance_policy": ["ATE insurance", "adverse costs", "insurance policy"],
            "medical_report": ["medical report", "medical evidence"],
            "expert_report": ["expert report", "expert opinion", "expert witness"],
            "settlement_offer": ["settlement", "part 36", "offer to settle"]
        }
        
        documentation_received = []
        text_lower = text.lower()
        for doc_type, keywords in doc_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                documentation_received.append(doc_type)
        
        extracted_data["documentation_received"] = documentation_received
        
        return extracted_data
    
    def _normalize_date(self, date_str: str) -> str:
        """Normalize date to ISO format"""
        formats = [
            "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y", "%d-%m-%y",
            "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M:%S", "%d-%m-%Y %H:%M:%S"
        ]

        for fmt in formats:
            try:
                dt = datetime.strptime(date_str, fmt)
                return dt.strftime("%Y-%m-%d")
            except ValueError:
                continue

        # Try to extract just the date part if time is present
        if isinstance(date_str, str) and " " in date_str:
            date_only = date_str.split(" ")[0]
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"]:
                try:
                    dt = datetime.strptime(date_only, fmt)
                    return dt.strftime("%Y-%m-%d")
                except ValueError:
                    continue

        return date_str
    
    def extract_financial_summary(self, text: str) -> Dict[str, Any]:
        """Extract detailed financial information"""
        financial_data = {
            "costs_incurred": [],
            "projected_costs": [],
            "disbursements": []
        }
        
        # Extract cost items
        cost_pattern = re.compile(r"([\w\s]+):\s*£([\d,]+(?:\.\d{2})?)")
        matches = cost_pattern.findall(text)
        
        for item, amount in matches:
            item_clean = item.strip().lower()
            amount_clean = float(amount.replace(",", ""))
            
            if "solicitor" in item_clean or "legal fee" in item_clean:
                financial_data["costs_incurred"].append({
                    "category": "legal_fees",
                    "description": item.strip(),
                    "amount": amount_clean
                })
            elif "court fee" in item_clean:
                financial_data["costs_incurred"].append({
                    "category": "court_fees",
                    "description": item.strip(),
                    "amount": amount_clean
                })
            elif "expert" in item_clean or "witness" in item_clean:
                financial_data["costs_incurred"].append({
                    "category": "expert_fees",
                    "description": item.strip(),
                    "amount": amount_clean
                })
        
        return financial_data
    
    def extract_case_milestones(self, text: str) -> List[Dict[str, str]]:
        """Extract case milestones and timeline"""
        milestones = []
        
        milestone_keywords = {
            "claim_filed": ["claim filed", "claim issued", "proceedings commenced"],
            "defense_received": ["defense filed", "defense received", "defense served"],
            "disclosure": ["disclosure", "disclosure exchange"],
            "expert_exchange": ["expert report", "expert evidence exchanged"],
            "trial_date": ["trial date", "hearing listed", "trial listed"],
            "settlement": ["settled", "settlement reached", "agreement reached"]
        }
        
        text_lower = text.lower()
        for milestone_type, keywords in milestone_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    # Try to find a date near the keyword
                    context_start = max(0, text_lower.index(keyword) - 50)
                    context_end = min(len(text), text_lower.index(keyword) + 100)
                    context = text[context_start:context_end]
                    
                    date_match = self.extraction_patterns["date"].search(context)
                    milestone_date = None
                    if date_match:
                        date_str = date_match.group(1) if date_match.group(1) else date_match.group(2)
                        milestone_date = self._normalize_date(date_str)
                    
                    milestones.append({
                        "milestone": milestone_type,
                        "date": milestone_date or "Date not specified",
                        "description": keyword.title()
                    })
                    break
        
        return milestones
    
    def extract_milberg_portfolio_summary(self, file_path: str) -> Dict[str, Any]:
        """
        Extract portfolio summary from Milberg monthly report
        """
        try:
            df = pd.read_excel(file_path, sheet_name='Portfolio Summary', engine='openpyxl')

            if len(df) > 0:
                row = df.iloc[0]
                summary = {
                    'report_date': str(row.get('Report Date', '')),
                    'total_bundles_funded': int(row.get('Total Bundles Funded', 0)) if pd.notna(row.get('Total Bundles Funded')) else 0,
                    'total_claimants': int(row.get('Total Claimants', 0)) if pd.notna(row.get('Total Claimants')) else 0,
                    'total_claims_submitted': int(row.get('Total Claims Submitted', 0)) if pd.notna(row.get('Total Claims Submitted')) else 0,
                    'total_funding_provided': float(row.get('Total Funding Provided (£)', 0)) if pd.notna(row.get('Total Funding Provided (£)')) else 0,
                    'outstanding_funding': float(row.get('Outstanding Funding (£)', 0)) if pd.notna(row.get('Outstanding Funding (£)')) else 0,
                    'total_dba_proceeds': float(row.get('Total DBA Proceeds (£)', 0)) if pd.notna(row.get('Total DBA Proceeds (£)')) else 0,
                    'funder_total_share': float(row.get('Funder Total Share (£)', 0)) if pd.notna(row.get('Funder Total Share (£)')) else 0,
                    'milberg_total_share': float(row.get('Milberg Total Share (£)', 0)) if pd.notna(row.get('Milberg Total Share (£)')) else 0,
                    'average_redress': float(row.get('Average Redress (£)', 0)) if pd.notna(row.get('Average Redress (£)')) else 0,
                    'redress_calculated_count': int(row.get('Redress Calculated (count)', 0)) if pd.notna(row.get('Redress Calculated (count)')) else 0,
                    'paid_to_claimants_count': int(row.get('Paid to Claimants (count)', 0)) if pd.notna(row.get('Paid to Claimants (count)')) else 0,
                    'respondents_engaged_count': int(row.get('Respondents Engaged (count)', 0)) if pd.notna(row.get('Respondents Engaged (count)')) else 0,
                    'remarks': str(row.get('Remarks', ''))
                }
                return summary
        except Exception as e:
            print(f"Error extracting portfolio summary: {str(e)}")

        return {}

    def extract_milberg_bundle_tracker(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract bundle tracker data from Milberg monthly report
        """
        try:
            df = pd.read_excel(file_path, sheet_name='Bundle Tracker', engine='openpyxl')

            bundles = []
            for idx, row in df.iterrows():
                bundle_id = row.get('Bundle ID')
                if pd.notna(bundle_id) and str(bundle_id).strip():
                    bundle = {
                        'bundle_id': str(bundle_id),
                        'bundle_name': str(row.get('Bundle Name', '')),
                        'funding_date': str(row.get('Funding Date', '')) if pd.notna(row.get('Funding Date')) else '',
                        'claimants_in_bundle': int(row.get('Claimants in Bundle', 0)) if pd.notna(row.get('Claimants in Bundle')) else 0,
                        'claims_submitted': int(row.get('Claims Submitted to Scheme', 0)) if pd.notna(row.get('Claims Submitted to Scheme')) else 0,
                        'redress_calculated_count': int(row.get('Redress Calculated (count)', 0)) if pd.notna(row.get('Redress Calculated (count)')) else 0,
                        'paid_to_claimants_count': int(row.get('Paid to Claimants (count)', 0)) if pd.notna(row.get('Paid to Claimants (count)')) else 0,
                        'total_action_costs': float(row.get('Total Action Costs (£)', 0)) if pd.notna(row.get('Total Action Costs (£)')) else 0,
                        'funding_drawn': float(row.get('Funding Drawn (£)', 0)) if pd.notna(row.get('Funding Drawn (£)')) else 0,
                        'cumulative_funding': float(row.get('Cumulative Funding (£)', 0)) if pd.notna(row.get('Cumulative Funding (£)')) else 0,
                        'dba_proceeds_received': float(row.get('DBA Proceeds Received (£)', 0)) if pd.notna(row.get('DBA Proceeds Received (£)')) else 0,
                        'projected_recoveries': float(row.get('Projected Recoveries (£)', 0)) if pd.notna(row.get('Projected Recoveries (£)')) else 0,
                        'funder_share': float(row.get('Funder Share (£)', 0)) if pd.notna(row.get('Funder Share (£)')) else 0,
                        'milberg_share': float(row.get('Milberg Share (£)', 0)) if pd.notna(row.get('Milberg Share (£)')) else 0,
                        'current_status': str(row.get('Current Status', '')) if pd.notna(row.get('Current Status')) else '',
                        'notes': str(row.get('Notes', '')) if pd.notna(row.get('Notes')) else ''
                    }
                    bundles.append(bundle)

            return bundles
        except Exception as e:
            print(f"Error extracting bundle tracker: {str(e)}")
            return []

    def extract_milberg_bundle_claims(self, file_path: str, bundle_sheet_name: str) -> List[Dict[str, Any]]:
        """
        Extract individual claim data from a specific bundle sheet
        """
        try:
            df = pd.read_excel(file_path, sheet_name=bundle_sheet_name, engine='openpyxl')

            claims = []
            for idx, row in df.iterrows():
                claimant_id = row.get('Claimant ID')
                # Skip total/summary rows and empty rows
                if pd.notna(claimant_id) and str(claimant_id).strip():
                    claimant_id_str = str(claimant_id).strip().upper()
                    # Skip rows that are totals or summaries
                    if claimant_id_str in ['TOTAL', 'TOTALS', 'SUMMARY', 'SUBTOTAL']:
                        continue
                    if claimant_id_str.startswith('TOTAL') or claimant_id_str.startswith('SUMMARY'):
                        continue
                    # Robust commission extraction
                    def parse_commission(val, pct=False):
                        if pd.isna(val):
                            return 0.0
                        if isinstance(val, str):
                            val = val.strip().replace('£', '').replace('%', '').replace(',', '').replace(' ', '')
                            try:
                                return float(val)
                            except ValueError:
                                return 0.0
                        try:
                            return float(val)
                        except Exception:
                            return 0.0


                    # Try multiple possible column names for commission fields
                    commission_pct = None
                    commission_amt = None
                    submission_date = None
                    pct_names = ['Commission % of Cost of Credit', 'Commission Percentage (%)', 'Commission %', 'Commission Percentage', 'I']
                    amt_names = ['Commission Amount (£)', 'Commission Amount', 'Commission (£)', 'J']
                    sub_names = ['Submission Date', 'N']
                    for name in pct_names:
                        if name in df.columns:
                            commission_pct = row[name]
                            break
                    for name in amt_names:
                        if name in df.columns:
                            commission_amt = row[name]
                            break
                    for name in sub_names:
                        if name in df.columns:
                            submission_date = row[name]
                            break

                    commission_pct_val = parse_commission(commission_pct, pct=True)
                    commission_amt_val = parse_commission(commission_amt)
                    submission_date_val = str(submission_date).strip() if pd.notna(submission_date) else ''

                    # Convert commission percentage from decimal to percentage if needed
                    # If the value is between 0 and 1, it's likely a decimal (e.g., 0.5 = 50%)
                    if commission_pct_val > 0 and commission_pct_val < 1:
                        commission_pct_val = commission_pct_val * 100

                    claim = {
                        'claimant_id': str(claimant_id),
                        'claim_id': str(claimant_id),
                        'respondent': str(row.get('Respondent', '')) if pd.notna(row.get('Respondent')) else '',
                        'defendant': str(row.get('Respondent', '')) if pd.notna(row.get('Respondent')) else '',
                        'agreement_date': str(row.get('Agreement Date', '')) if pd.notna(row.get('Agreement Date')) else '',
                        'product_type': str(row.get('Product Type', '')) if pd.notna(row.get('Product Type')) else '',
                        'loan_amount': float(row.get('Loan Amount (£)', 0)) if pd.notna(row.get('Loan Amount (£)')) else 0,
                        'total_cost_of_credit': float(row.get('Total Cost of Credit (£)', 0)) if pd.notna(row.get('Total Cost of Credit (£)')) else 0,
                        'commission_pct_of_cost': commission_pct_val,
                        'commission_amount': commission_amt_val,
                        'submission_date': submission_date_val,
                        'decision_status': str(row.get('Decision Status', '')) if pd.notna(row.get('Decision Status')) else '',
                        'status': str(row.get('Decision Status', 'in_progress')).lower().replace(' ', '_') if pd.notna(row.get('Decision Status')) else 'in_progress',
                        'redress_calculated': float(row.get('Redress Calculated (£)', 0)) if pd.notna(row.get('Redress Calculated (£)')) else 0,
                        'claim_amount': float(row.get('Redress Calculated (£)', 0)) if pd.notna(row.get('Redress Calculated (£)')) else 0,
                        'claimant_payout': float(row.get('Claimant Payout (£)', 0)) if pd.notna(row.get('Claimant Payout (£)')) else 0,
                        'dba_proceeds': float(row.get('DBA Proceeds (£)', 0)) if pd.notna(row.get('DBA Proceeds (£)')) else 0,
                        'funder_share': float(row.get('Funder Share (£)', 0)) if pd.notna(row.get('Funder Share (£)')) else 0,
                        'bundle_sheet': bundle_sheet_name,
                        'law_firm': 'Milberg'
                    }
                    claims.append(claim)

            return claims
        except Exception as e:
            print(f"Error extracting claims from {bundle_sheet_name}: {str(e)}")
            return []

    def extract_from_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract structured data from Excel files (Milberg monthly reports)
        Returns a list of claim records extracted from the spreadsheet
        """
        try:
            # Check if this is a Milberg report by looking for specific sheets
            xl = pd.ExcelFile(file_path)

            if 'Bundle Tracker' in xl.sheet_names and 'Portfolio Summary' in xl.sheet_names:
                # This is a Milberg report - extract all claims from bundle sheets
                all_claims = []

                for sheet_name in xl.sheet_names:
                    if sheet_name.startswith('Bundle_') and not sheet_name == 'Template_Bundle':
                        bundle_claims = self.extract_milberg_bundle_claims(file_path, sheet_name)
                        all_claims.extend(bundle_claims)

                return all_claims

            else:
                # Generic Excel processing
                df = pd.read_excel(file_path, sheet_name=0, engine='openpyxl')

                records = []
                column_mapping = {
                    'claim_id': ['claim id', 'claim_id', 'reference', 'ref', 'claim reference', 'case number', 'claimant id'],
                    'claimant_name': ['claimant', 'claimant name', 'client', 'client name'],
                    'defendant': ['defendant', 'defendant name', 'lender', 'finance company', 'respondent'],
                    'claim_amount': ['claim amount', 'claim value', 'total claim', 'amount', 'redress calculated'],
                    'funded_amount': ['funded', 'funded amount', 'funding', 'advanced'],
                    'status': ['status', 'case status', 'state', 'decision status'],
                    'law_firm': ['law firm', 'solicitor', 'firm', 'lawyer'],
                    'claim_date': ['date', 'claim date', 'filed date', 'start date', 'agreement date'],
                }

                for idx, row in df.iterrows():
                    record = {}

                    for target_field, possible_names in column_mapping.items():
                        for col_name in df.columns:
                            if col_name.lower().strip() in possible_names:
                                value = row[col_name]
                                if pd.notna(value):
                                    if target_field in ['claim_amount', 'funded_amount']:
                                        try:
                                            if isinstance(value, str):
                                                value = value.replace('£', '').replace('$', '').replace(',', '').strip()
                                            record[target_field] = float(value)
                                        except (ValueError, TypeError):
                                            pass
                                    elif target_field == 'claim_date':
                                        try:
                                            if isinstance(value, datetime):
                                                record[target_field] = value.strftime('%Y-%m-%d')
                                            else:
                                                record[target_field] = str(value)
                                        except:
                                            record[target_field] = str(value)
                                    elif target_field == 'status':
                                        status_str = str(value).lower().strip()
                                        record[target_field] = status_str.replace(' ', '_')
                                    else:
                                        record[target_field] = str(value).strip()
                                break

                    if record.get('claim_id') or record.get('claimant_name'):
                        if 'status' not in record:
                            record['status'] = 'in_progress'
                        if 'law_firm' not in record:
                            record['law_firm'] = 'Unknown'

                        record['documentation_received'] = []
                        record['processed_at'] = datetime.now().isoformat()
                        record['source_file'] = file_path
                        record['last_update'] = datetime.now().strftime('%Y-%m-%d')

                        records.append(record)

                return records

        except Exception as e:
            print(f"Error processing Excel file: {str(e)}")
            import traceback
            traceback.print_exc()
            return []

    def extract_from_word(self, file_path: str) -> Dict[str, Any]:
        """
        Extract structured data from Word documents (e.g., Milberg monthly report in .docx format)
        Uses OpenAI to intelligently parse and structure the document content
        """
        try:
            doc = Document(file_path)

            # Extract all text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Extract text from tables
            tables_data = []
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_text.append(row_data)
                if table_text:
                    tables_data.append(table_text)
                    # Convert table to readable text format
                    tables_text.append('\n'.join([' | '.join(row) for row in table_text]))

            # Combine all text
            combined_text = '\n'.join(full_text)
            all_tables_text = '\n\n--- TABLE ---\n\n'.join(tables_text)
            full_document_text = f"{combined_text}\n\n{'='*50}\nTABLES:\n{'='*50}\n\n{all_tables_text}"

            print(f"[Info] Extracted {len(tables_data)} tables from Word document")
            print(f"[Info] Sending to OpenAI for intelligent parsing...")

            # Use OpenAI to intelligently parse the document
            prompt = f"""You are a legal document analyst specializing in PCP (Personal Contract Purchase) claims reports.

Analyze the following Milberg monthly report document and extract ALL claim information in a structured JSON format.

DOCUMENT CONTENT:
{full_document_text[:15000]}

Your task:
1. Identify ALL individual claims mentioned in the document (from tables or text)
2. For EACH claim, extract the following fields (use null if not found):
   - claim_id or claimant_id (any reference number)
   - defendant or respondent (lender name)
   - claim_amount or redress_calculated (monetary value in £)
   - funded_amount or funder_share (monetary value in £)
   - status or decision_status
   - agreement_date
   - product_type (e.g., PCP, HP, Conditional Sale)
   - loan_amount
   - commission_pct_of_cost (percentage)
   - submission_date

3. Also extract portfolio summary information:
   - total_bundles_funded
   - total_claimants
   - total_claims_submitted
   - total_funding_provided
   - report_date

Return a JSON object with this structure:
{{
  "report_type": "Milberg Monthly Report",
  "report_date": "YYYY-MM-DD",
  "portfolio_summary": {{
    "total_bundles_funded": 0,
    "total_claimants": 0,
    "total_claims_submitted": 0,
    "total_funding_provided": 0.0,
    "total_dba_proceeds": 0.0,
    "funder_total_share": 0.0
  }},
  "claims": [
    {{
      "claim_id": "...",
      "claimant_id": "...",
      "defendant": "...",
      "claim_amount": 0.0,
      "funded_amount": 0.0,
      "status": "...",
      "agreement_date": "...",
      "product_type": "...",
      "loan_amount": 0.0,
      "commission_pct_of_cost": 0.0,
      "submission_date": "...",
      "law_firm": "Milberg"
    }}
  ],
  "bundle_tracker": []
}}

IMPORTANT:
- Extract ALL claims you can find
- Use numeric values for amounts (not strings)
- Use "Milberg" as the law_firm for all claims
- Return ONLY valid JSON, no additional text"""

            # Call OpenAI
            llm_response = self.call_openai_llm(prompt, model="gpt-4o", max_tokens=4000)

            # Parse JSON response
            try:
                # Clean response (remove markdown code blocks if present)
                cleaned_response = llm_response.strip()
                if cleaned_response.startswith('```'):
                    # Remove markdown code blocks
                    lines = cleaned_response.split('\n')
                    cleaned_response = '\n'.join([l for l in lines if not l.startswith('```')])

                extracted_data = json.loads(cleaned_response)

                # Add metadata
                extracted_data['processed_at'] = datetime.now().isoformat()
                extracted_data['source_file'] = file_path
                extracted_data['last_update'] = datetime.now().strftime('%Y-%m-%d')
                extracted_data['tables_found'] = len(tables_data)
                extracted_data['extraction_method'] = 'openai_llm'
                extracted_data['total_claims'] = len(extracted_data.get('claims', []))

                print(f"[Success] OpenAI extracted {extracted_data['total_claims']} claims")

                return extracted_data

            except json.JSONDecodeError as e:
                print(f"[Warning] Could not parse OpenAI response as JSON: {e}")
                print(f"[Debug] LLM Response: {llm_response[:500]}")

                # Fallback to basic extraction
                return self._fallback_word_extraction(combined_text, tables_data, file_path)

        except Exception as e:
            print(f"Error extracting from Word document: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'error': str(e),
                'report_type': 'Word Document Report',
                'source_file': file_path,
                'processed_at': datetime.now().isoformat()
            }

    def _fallback_word_extraction(self, combined_text: str, tables_data: List, file_path: str) -> Dict[str, Any]:
        """
        Fallback method for Word extraction if OpenAI parsing fails
        """
        print("[Info] Using fallback extraction method")

        # Use existing text extraction logic
        extracted_data = self.extract_from_text(combined_text)

        # Add metadata
        extracted_data['report_type'] = 'Word Document Report'
        extracted_data['processed_at'] = datetime.now().isoformat()
        extracted_data['source_file'] = file_path
        extracted_data['last_update'] = datetime.now().strftime('%Y-%m-%d')
        extracted_data['tables_found'] = len(tables_data)
        extracted_data['extraction_method'] = 'fallback'

        # Try to extract structured data from tables if present
        claims = []
        if tables_data:
            # Process tables to find claim information
            for table_data in tables_data:
                if len(table_data) > 1:  # Has header and data rows
                    header_row = [h.lower().strip() for h in table_data[0]]

                    # Check if this looks like a claims table
                    claim_indicators = ['claimant', 'claim', 'bundle', 'redress', 'respondent']
                    if any(indicator in ' '.join(header_row) for indicator in claim_indicators):
                        # Process data rows
                        for data_row in table_data[1:]:
                            claim = self._extract_claim_from_table_row(header_row, data_row)
                            if claim:
                                claims.append(claim)

        if claims:
            extracted_data['claims'] = claims
            extracted_data['total_claims'] = len(claims)

            # Calculate summary statistics
            total_claim_value = sum(c.get('claim_amount', 0) for c in claims)
            total_funded = sum(c.get('funded_amount', 0) for c in claims)

            extracted_data['portfolio_summary'] = {
                'total_claims': len(claims),
                'total_claim_value': total_claim_value,
                'total_funded': total_funded,
                'extracted_from': 'word_tables'
            }
        else:
            # Single claim document
            extracted_data['claims'] = [extracted_data.copy()]
            extracted_data['total_claims'] = 1

        return extracted_data

    def _extract_claim_from_table_row(self, headers: List[str], row_data: List[str]) -> Optional[Dict[str, Any]]:
        """
        Extract claim information from a table row based on headers
        """
        try:
            claim = {
                'law_firm': 'Milberg',
                'status': 'in_progress'
            }

            # Map common column names to our fields
            column_mappings = {
                'claimant_id': ['claimant id', 'claimant', 'client id', 'id'],
                'claim_id': ['claim id', 'reference', 'ref', 'claimant id'],
                'defendant': ['defendant', 'respondent', 'lender'],
                'claim_amount': ['claim amount', 'redress calculated', 'redress', 'amount'],
                'funded_amount': ['funder share', 'funding', 'funded'],
                'agreement_date': ['agreement date', 'date', 'signed date'],
                'status': ['status', 'decision status', 'state'],
                'product_type': ['product type', 'product', 'loan type'],
                'loan_amount': ['loan amount', 'principal'],
                'commission_pct_of_cost': ['commission %', 'commission', 'commission percentage'],
            }

            for field, possible_headers in column_mappings.items():
                for idx, header in enumerate(headers):
                    if idx < len(row_data) and any(ph in header for ph in possible_headers):
                        value = row_data[idx].strip()
                        if value:
                            # Type conversion
                            if field in ['claim_amount', 'funded_amount', 'loan_amount']:
                                try:
                                    # Remove currency symbols and commas
                                    clean_val = value.replace('£', '').replace('$', '').replace(',', '').strip()
                                    claim[field] = float(clean_val) if clean_val else 0
                                except ValueError:
                                    pass
                            elif field == 'commission_pct_of_cost':
                                try:
                                    clean_val = value.replace('%', '').replace(',', '').strip()
                                    pct = float(clean_val) if clean_val else 0
                                    # Convert decimal to percentage if needed
                                    if pct > 0 and pct < 1:
                                        pct = pct * 100
                                    claim[field] = pct
                                except ValueError:
                                    pass
                            else:
                                claim[field] = value
                        break

            # Set claim_id if not set but claimant_id exists
            if 'claim_id' not in claim and 'claimant_id' in claim:
                claim['claim_id'] = claim['claimant_id']

            # Only return if we have at least a claim_id
            if claim.get('claim_id'):
                return claim

            return None

        except Exception as e:
            print(f"Error extracting claim from table row: {str(e)}")
            return None

    def extract_from_excel_detailed(self, file_path: str) -> Dict[str, Any]:
        """
        Extract detailed information from Excel including summary statistics
        """
        # Check if this is a Milberg report
        try:
            xl = pd.ExcelFile(file_path)
            is_milberg = 'Bundle Tracker' in xl.sheet_names and 'Portfolio Summary' in xl.sheet_names
        except:
            is_milberg = False

        if is_milberg:
            # Extract Milberg-specific data
            portfolio_summary = self.extract_milberg_portfolio_summary(file_path)
            bundle_tracker = self.extract_milberg_bundle_tracker(file_path)
            claims = self.extract_from_excel(file_path)

            summary = {
                "report_type": "Milberg Monthly Report",
                "portfolio_summary": portfolio_summary,
                "bundle_tracker": bundle_tracker,
                "total_claims": len(claims),
                "total_funded": portfolio_summary.get('total_funding_provided', 0),
                "total_claim_value": sum(r.get('claim_amount', 0) for r in claims),
                "total_bundles": len(bundle_tracker),
                "claims": claims,
                "processed_at": datetime.now().isoformat(),
                "source_file": file_path
            }

            # Status breakdown
            status_counts = {}
            for record in claims:
                status = record.get('status', 'unknown')
                status_counts[status] = status_counts.get(status, 0) + 1

            summary["status_breakdown"] = status_counts

        else:
            # Generic Excel processing
            records = self.extract_from_excel(file_path)

            summary = {
                "report_type": "Generic Excel Report",
                "total_claims": len(records),
                "total_funded": sum(r.get('funded_amount', 0) for r in records),
                "total_claim_value": sum(r.get('claim_amount', 0) for r in records),
                "claims": records,
                "processed_at": datetime.now().isoformat(),
                "source_file": file_path
            }

            # Status breakdown
            status_counts = {}
            for record in records:
                status = record.get('status', 'unknown')
                status_counts[status] = status_counts.get(status, 0) + 1

            summary["status_breakdown"] = status_counts

        return summary

    def process_law_firm_report(self, file_path: str) -> Dict[str, Any]:
        """
        Main entry point for processing law firm reports
        Supports TXT, XLSX, XLS, DOCX files
        """
        # Determine file type
        file_extension = file_path.lower().split('.')[-1]

        if file_extension in ['xlsx', 'xls']:
            # Process Excel file
            excel_data = self.extract_from_excel_detailed(file_path)
            return excel_data

        elif file_extension == 'docx':
            # Process Word document
            word_data = self.extract_from_word(file_path)
            return word_data

        else:
            # Process text file
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()

            # Extract core data
            extracted_data = self.extract_from_text(text)

            # Extract financial details
            financial_data = self.extract_financial_summary(text)
            extracted_data["financial_breakdown"] = financial_data

            # Extract milestones
            milestones = self.extract_case_milestones(text)
            extracted_data["case_milestones"] = milestones

            # Add metadata
            extracted_data["processed_at"] = datetime.now().isoformat()
            extracted_data["source_file"] = file_path
            extracted_data["last_update"] = datetime.now().strftime("%Y-%m-%d")

            return extracted_data
    
    def validate_extracted_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate extracted data and flag missing required fields
        """
        required_fields = [
            "claim_id", "claimant_name", "defendant", 
            "claim_amount", "funded_amount", "status", "law_firm"
        ]
        
        validation_result = {
            "is_valid": True,
            "missing_fields": [],
            "warnings": []
        }
        
        for field in required_fields:
            if field not in data or not data[field]:
                validation_result["missing_fields"].append(field)
                validation_result["is_valid"] = False
        
        # Check data types and ranges
        if "claim_amount" in data and data["claim_amount"] <= 0:
            validation_result["warnings"].append("Claim amount must be positive")
        
        if "funded_amount" in data and "claim_amount" in data:
            if data["funded_amount"] > data["claim_amount"]:
                validation_result["warnings"].append("Funded amount exceeds claim amount")
        
        return validation_result


def demo_document_processing():
    """Demonstrate document processing capabilities"""
    
    # Sample law firm report text
    sample_report = """
    LEGAL PARTNERS LLP
    Claim Progress Report
    
    Claim Reference: PCP/2024/5678
    Date: 15/11/2024
    
    Claimant: Sarah Johnson
    Defendant: Auto Finance Solutions Limited
    
    FINANCIAL SUMMARY
    Claim Amount: £95,500.00
    Funded Amount: £42,000.00
    
    Status: In Progress
    
    COSTS BREAKDOWN
    Solicitor Fees: £18,500.00
    Court Fees: £2,500.00
    Expert Witness Fees: £5,000.00
    ATE Insurance Premium: £8,000.00
    
    DOCUMENTATION
    - Claim Form completed and filed
    - Litigation Funding Agreement signed
    - ATE Insurance policy in place
    - Expert Report commissioned
    
    CASE MILESTONES
    Claim Filed: 10/09/2024
    Defense Received: 25/09/2024
    Disclosure deadline: 15/12/2024
    Trial Date: 15/03/2025
    
    The case is progressing well with no significant delays anticipated.
    """
    
    processor = DocumentProcessor()
    
    # Extract data
    print("="*60)
    print("DOCUMENT PROCESSING DEMO")
    print("="*60)
    
    extracted = processor.extract_from_text(sample_report)
    print("\nExtracted Data:")
    print(json.dumps(extracted, indent=2))
    
    # Validate
    validation = processor.validate_extracted_data(extracted)
    print("\nValidation Result:")
    print(json.dumps(validation, indent=2))
    
    # Extract milestones
    milestones = processor.extract_case_milestones(sample_report)
    print("\nCase Milestones:")
    print(json.dumps(milestones, indent=2))


if __name__ == "__main__":
    demo_document_processing()
