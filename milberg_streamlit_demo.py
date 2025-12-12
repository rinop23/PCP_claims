"""
Enhanced PCP Claims Analysis Dashboard
Multi-Agent AI System with Intelligent Visualizations
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import os
import hashlib
from intelligent_agents import generate_full_investor_report
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64

# Page config
st.set_page_config(
    page_title="PCP Claims Investor Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f77b4;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .agent-status {
        padding: 0.5rem;
        border-radius: 0.3rem;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Authentication
def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

USERS = {
    "admin": hash_password("Admin123!"),
    "walter": hash_password("Walter123!"),
    "dirk": hash_password("Dirk123!"),
    "eda": hash_password("Eda123!")
}

def check_login(username: str, password: str) -> bool:
    if username in USERS:
        return USERS[username] == hash_password(password)
    return False

# Initialize session state
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'username' not in st.session_state:
    st.session_state.username = None
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None

# Helper functions for visualizations
def create_portfolio_overview_chart(portfolio_metrics):
    """Create portfolio overview metrics chart"""
    fig = go.Figure()

    categories = ['Submitted', 'Successful', 'Rejected', 'Pending']
    values = [
        portfolio_metrics.get('claims_submitted', 0),
        portfolio_metrics.get('claims_successful', 0),
        portfolio_metrics.get('claims_rejected', 0),
        portfolio_metrics.get('claims_pending', 0)
    ]

    fig.add_trace(go.Bar(
        x=categories,
        y=values,
        text=values,
        textposition='auto',
        marker_color=['#3498db', '#2ecc71', '#e74c3c', '#f39c12']
    ))

    fig.update_layout(
        title="Claims by Status",
        xaxis_title="Status",
        yaxis_title="Number of Claims",
        height=400,
        showlegend=False
    )

    return fig

def create_lender_distribution_chart(lender_data):
    """Create lender distribution pie chart for top 10"""
    df = pd.DataFrame(lender_data[:10])

    fig = px.pie(
        df,
        values='num_claims',
        names='lender',
        title='Top 10 Lenders by Claim Volume',
        hole=0.4
    )

    fig.update_traces(textposition='inside', textinfo='percent+label')
    fig.update_layout(height=500)

    return fig

def create_pipeline_funnel(pipeline_data):
    """Create pipeline funnel chart"""
    stages = ['Awaiting DSAR', 'Pending Submission', 'Under Review', 'Settlement Offered', 'Paid']
    counts = [
        pipeline_data.get('awaiting_dsar', {}).get('count', 0),
        pipeline_data.get('pending_submission', {}).get('count', 0),
        pipeline_data.get('under_review', {}).get('count', 0),
        pipeline_data.get('settlement_offered', {}).get('count', 0),
        pipeline_data.get('paid', {}).get('count', 0)
    ]

    fig = go.Figure(go.Funnel(
        y=stages,
        x=counts,
        textinfo="value+percent initial",
        marker={"color": ["#3498db", "#2ecc71", "#f39c12", "#9b59b6", "#1abc9c"]}
    ))

    fig.update_layout(
        title="Claims Pipeline Funnel",
        height=500
    )

    return fig

def create_financial_waterfall(financial_data):
    """Create waterfall chart for financial breakdown"""

    total_settlement = financial_data.get('total_settlements', 0)
    dba_proceeds = financial_data.get('dba_proceeds_expected', 0)
    costs = financial_data.get('total_costs_incurred', 0)
    net_proceeds = financial_data.get('net_proceeds_after_costs', 0)

    fig = go.Figure(go.Waterfall(
        name="Financial Flow",
        orientation="v",
        measure=["absolute", "relative", "relative", "total"],
        x=["Total Settlement", "DBA Fee (30%)", "Less: Costs", "Net Proceeds"],
        y=[total_settlement, dba_proceeds - total_settlement, -costs, net_proceeds],
        text=[f"£{total_settlement:,.0f}", f"£{dba_proceeds:,.0f}", f"-£{costs:,.0f}", f"£{net_proceeds:,.0f}"],
        textposition="outside",
        connector={"line": {"color": "rgb(63, 63, 63)"}},
    ))

    fig.update_layout(
        title="Financial Waterfall: Settlement to Net Proceeds",
        height=500,
        showlegend=False
    )

    return fig

def create_profit_split_chart(financial_data, profit_rules):
    """Create profit split pie chart"""

    funder_return = financial_data.get('funder_expected_return', 0)
    firm_return = financial_data.get('firm_expected_return', 0)

    funder_pct = profit_rules.get('profit_split', {}).get('funder_percentage', 80)
    firm_pct = profit_rules.get('profit_split', {}).get('law_firm_percentage', 20)

    fig = go.Figure(data=[go.Pie(
        labels=[f'Funder ({funder_pct}%)', f'Law Firm ({firm_pct}%)'],
        values=[funder_return, firm_return],
        hole=0.3,
        marker_colors=['#3498db', '#2ecc71']
    )])

    fig.update_traces(textposition='inside', textinfo='percent+label+value')
    fig.update_layout(
        title="Profit Distribution (After Costs)",
        height=400,
        annotations=[dict(text='Net Proceeds<br>Split', x=0.5, y=0.5, font_size=12, showarrow=False)]
    )

    return fig

def create_cost_breakdown_chart(financial_metrics):
    """Create cost breakdown chart"""

    costs = {
        'Acquisition': financial_metrics.get('acquisition_cost', 0),
        'Submission': financial_metrics.get('submission_cost', 0),
        'Processing': financial_metrics.get('processing_cost', 0),
        'Legal': financial_metrics.get('legal_cost', 0)
    }

    fig = px.bar(
        x=list(costs.keys()),
        y=list(costs.values()),
        title="Cost Breakdown by Category",
        labels={'x': 'Cost Category', 'y': 'Amount (£)'},
        text=[f"£{v:,.0f}" for v in costs.values()]
    )

    fig.update_traces(textposition='outside', marker_color='#e74c3c')
    fig.update_layout(height=400)

    return fig

def create_lender_value_chart(lender_data):
    """Create lender value horizontal bar chart"""
    df = pd.DataFrame(lender_data[:15])
    df = df.sort_values('estimated_value', ascending=True)

    fig = px.bar(
        df,
        x='estimated_value',
        y='lender',
        orientation='h',
        title='Top 15 Lenders by Portfolio Value',
        labels={'estimated_value': 'Portfolio Value (£)', 'lender': 'Lender'},
        text=[f"£{v:,.0f}" for v in df['estimated_value']]
    )

    fig.update_traces(textposition='outside', marker_color='#9b59b6')
    fig.update_layout(height=600, yaxis={'categoryorder':'total ascending'})

    return fig

def generate_word_report(report_markdown, monthly_data, investor_report):
    """Generate Word document report"""
    doc = Document()

    # Title
    title = doc.add_heading('Monthly Investor Report', 0)
    title.alignment = 1  # Center

    doc.add_heading(f"Reporting Period: {monthly_data.get('reporting_period', 'N/A')}", level=2)
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = investor_report.get('executive_summary', {})
    doc.add_paragraph(f"Portfolio Health Score: {exec_summary.get('portfolio_health_score', 0)}/100")
    doc.add_paragraph()

    # Key metrics
    doc.add_heading('Key Metrics', level=2)
    for metric in exec_summary.get('key_metrics_summary', []):
        doc.add_paragraph(metric, style='List Bullet')

    doc.add_page_break()

    # Portfolio Performance
    doc.add_heading('Portfolio Performance', level=1)
    perf = investor_report.get('portfolio_performance', {})

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    metrics = [
        ('Total Claims', f"{perf.get('total_claims', 0):,}"),
        ('Total Clients', f"{perf.get('total_clients', 0):,}"),
        ('Success Rate', f"{perf.get('success_rate', 0):.1f}%"),
        ('Average Settlement', f"£{perf.get('average_settlement', 0):,.2f}"),
        ('Portfolio Value', f"£{perf.get('total_portfolio_value', 0):,.2f}")
    ]

    for i, (label, value) in enumerate(metrics):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_page_break()

    # Financial Analysis
    doc.add_heading('Financial Analysis', level=1)
    fin = investor_report.get('financial_analysis', {})

    doc.add_heading('Revenue', level=2)
    doc.add_paragraph(f"Total Expected Settlements: £{fin.get('total_settlements', 0):,.2f}")
    doc.add_paragraph(f"DBA Proceeds (30%): £{fin.get('dba_proceeds_expected', 0):,.2f}")

    doc.add_heading('Profit Distribution', level=2)
    doc.add_paragraph(f"Funder Expected Return: £{fin.get('funder_expected_return', 0):,.2f}")
    doc.add_paragraph(f"Firm Expected Return: £{fin.get('firm_expected_return', 0):,.2f}")

    doc.add_heading('Performance Metrics', level=2)
    doc.add_paragraph(f"ROI Projection: {fin.get('roi_projection', 0):.1f}%")
    doc.add_paragraph(f"MOIC Projection: {fin.get('moic_projection', 0):.2f}x")

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# Login page
if not st.session_state.logged_in:
    st.markdown('<div class="main-header">🔐 PCP Claims Investor Dashboard</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">AI-Powered Multi-Agent Analysis System</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        with st.form("login_form"):
            st.subheader("Login")
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Login", use_container_width=True)

            if submit:
                if check_login(username, password):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.rerun()
                else:
                    st.error("Invalid username or password")

# Main dashboard
else:
    # Sidebar
    with st.sidebar:
        st.markdown("### 🤖 AI Agent System")
        st.markdown(f"**User:** {st.session_state.username}")

        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.username = None
            st.session_state.analysis_result = None
            st.rerun()

        st.markdown("---")

        st.markdown("### 📄 Active Agents")
        st.markdown("""
        1. 📄 **Priority Deed Agent**
           Reads profit distribution rules

        2. ⚖️ **FCA Compliance Agent**
           Validates claim compliance

        3. 📊 **Monthly Report Agent**
           Extracts Excel data

        4. 📝 **Investor Report Agent**
           Generates final report
        """)

        st.markdown("---")
        st.markdown("### ℹ️ System Info")
        st.markdown("**Model:** GPT-4o")
        st.markdown("**Analysis Time:** ~60-90 sec")

    # Header
    st.markdown('<div class="main-header">📊 PCP Claims Investor Dashboard</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Intelligent Multi-Agent Analysis System</div>', unsafe_allow_html=True)

    # File upload
    st.markdown("### 📤 Upload Monthly Report")
    uploaded_file = st.file_uploader(
        "Upload Milberg Monthly Report (Excel)",
        type=['xlsx', 'xls'],
        help="Upload the Monthly Summary Excel file to generate comprehensive investor report"
    )

    if uploaded_file:
        # Save uploaded file
        os.makedirs("uploads", exist_ok=True)
        temp_path = os.path.join("uploads", uploaded_file.name)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        col1, col2 = st.columns([3, 1])
        with col1:
            analyze_button = st.button("🤖 Generate Investor Report", type="primary", use_container_width=True)

        # Analyze button
        if analyze_button:
            with st.spinner("🤖 AI Agents are analyzing... This may take 60-90 seconds..."):

                # Progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()

                try:
                    status_text.text("📄 Priority Deed Agent: Reading profit distribution rules...")
                    progress_bar.progress(20)

                    result = generate_full_investor_report(temp_path)

                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()

                    st.session_state.analysis_result = result
                    st.success("✅ Investor report generated successfully!")
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Analysis failed: {str(e)}")
                    st.exception(e)

    # Display results
    if st.session_state.analysis_result:
        result = st.session_state.analysis_result
        monthly_data = result['monthly_data']
        investor_report = result['investor_report']
        profit_rules = result['profit_rules']

        st.markdown("---")

        # Key Metrics Dashboard
        st.markdown("### 📊 Key Performance Indicators")

        portfolio = monthly_data.get('portfolio_metrics', {})
        lenders = monthly_data.get('lender_distribution', [])
        financial = investor_report.get('financial_analysis', {})

        col1, col2, col3, col4, col5 = st.columns(5)

        with col1:
            st.metric(
                "Total Claims",
                f"{portfolio.get('unique_claims', 0):,}",
                help="Cumulative claims in portfolio"
            )

        with col2:
            st.metric(
                "Success Rate",
                f"{portfolio.get('success_rate', 0):.1f}%",
                f"{portfolio.get('claims_successful', 0)}/{portfolio.get('claims_submitted', 0)}",
                help="Claims successful vs submitted"
            )

        with col3:
            st.metric(
                "Portfolio Value",
                f"£{portfolio.get('total_settlement_value', 0)/1000:.0f}K",
                help="Total settlement value"
            )

        with col4:
            st.metric(
                "Expected Return",
                f"£{financial.get('funder_expected_return', 0)/1000:.0f}K",
                help="Funder's projected return"
            )

        with col5:
            health_score = investor_report.get('executive_summary', {}).get('portfolio_health_score', 0)
            st.metric(
                "Health Score",
                f"{health_score}/100",
                help="Overall portfolio health"
            )

        st.markdown("---")

        # Main content tabs
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
            "📋 Executive Report",
            "💰 Financial Analysis",
            "⚖️ Compliance",
            "📊 Portfolio Analytics",
            "📈 Visualizations",
            "📁 Raw Data"
        ])

        with tab1:
            st.markdown("## Monthly Investor Report")
            st.markdown(result['markdown_report'])

            # Download buttons
            col1, col2 = st.columns(2)

            with col1:
                st.download_button(
                    label="📥 Download Report (Markdown)",
                    data=result['markdown_report'],
                    file_name=f"investor_report_{monthly_data.get('reporting_period', 'report').replace('/', '_')}.md",
                    mime="text/markdown",
                    use_container_width=True
                )

            with col2:
                # Generate Word document
                word_doc = generate_word_report(
                    result['markdown_report'],
                    monthly_data,
                    investor_report
                )

                st.download_button(
                    label="📥 Download Report (Word)",
                    data=word_doc,
                    file_name=f"investor_report_{monthly_data.get('reporting_period', 'report').replace('/', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        with tab2:
            st.markdown("## Financial Performance Analysis")

            # Financial metrics
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("### Revenue & Proceeds")
                st.metric("Total Settlement Value", f"£{financial.get('total_settlements', 0):,.2f}")
                st.metric("DBA Proceeds (30%)", f"£{financial.get('dba_proceeds_expected', 0):,.2f}")
                st.metric("Net Proceeds (After Costs)", f"£{financial.get('net_proceeds_after_costs', 0):,.2f}")

            with col2:
                st.markdown("### Returns & Performance")
                st.metric("Funder Return", f"£{financial.get('funder_expected_return', 0):,.2f}")
                st.metric("Firm Return", f"£{financial.get('firm_expected_return', 0):,.2f}")
                st.metric("ROI Projection", f"{financial.get('roi_projection', 0):.1f}%")
                st.metric("MOIC Projection", f"{financial.get('moic_projection', 0):.2f}x")

            st.markdown("---")

            # Financial charts
            col1, col2 = st.columns(2)

            with col1:
                st.plotly_chart(
                    create_financial_waterfall(financial),
                    use_container_width=True
                )

            with col2:
                st.plotly_chart(
                    create_profit_split_chart(financial, profit_rules),
                    use_container_width=True
                )

            # Cost breakdown
            st.plotly_chart(
                create_cost_breakdown_chart(monthly_data.get('financial_metrics', {})),
                use_container_width=True
            )

        with tab3:
            st.markdown("## FCA Compliance Assessment")

            compliance = investor_report.get('compliance_assessment', {})

            # Compliance status
            status = compliance.get('fca_compliance_status', 'unknown')
            status_colors = {
                'compliant': '🟢',
                'review_required': '🟡',
                'at_risk': '🔴'
            }

            st.markdown(f"### {status_colors.get(status, '⚪')} Compliance Status: {status.upper().replace('_', ' ')}")

            col1, col2 = st.columns([2, 1])

            with col1:
                st.markdown("### Commission Analysis")
                st.info(compliance.get('commission_analysis', 'No analysis available'))

                st.markdown("### Required Actions")
                for action in compliance.get('compliance_actions_needed', []):
                    st.warning(f"⚠️ {action}")

            with col2:
                st.metric("Claims at Risk", compliance.get('claims_at_risk', 0))

                st.markdown("### Compliance Rules")
                st.json({
                    "FCA Thresholds": result.get('compliance_rules', {}).get('commission_thresholds', {}),
                    "Eligible Products": result.get('compliance_rules', {}).get('eligible_products', [])
                })

        with tab4:
            st.markdown("## Portfolio Composition & Analytics")

            # Lender concentration
            lender_conc = investor_report.get('lender_concentration', {})

            col1, col2, col3 = st.columns(3)

            with col1:
                st.metric("Total Lenders", lender_conc.get('total_lenders', 0))

            with col2:
                st.metric("Diversification Score", f"{lender_conc.get('diversification_score', 0)}/100")

            with col3:
                risk = lender_conc.get('concentration_risk', 'unknown')
                risk_colors = {'low': '🟢', 'medium': '🟡', 'high': '🔴'}
                st.metric("Concentration Risk", f"{risk_colors.get(risk, '⚪')} {risk.upper()}")

            st.markdown("---")

            # Charts
            col1, col2 = st.columns(2)

            with col1:
                st.plotly_chart(
                    create_lender_distribution_chart(lenders),
                    use_container_width=True
                )

            with col2:
                st.plotly_chart(
                    create_portfolio_overview_chart(portfolio),
                    use_container_width=True
                )

            # Top lenders table
            st.markdown("### Top 5 Lenders by Volume")
            top_lenders = lender_conc.get('top_5_lenders', [])
            if top_lenders:
                df_top = pd.DataFrame(top_lenders)
                st.dataframe(df_top, use_container_width=True, hide_index=True)

        with tab5:
            st.markdown("## Visual Analytics")

            # Pipeline funnel
            st.plotly_chart(
                create_pipeline_funnel(monthly_data.get('pipeline', {})),
                use_container_width=True
            )

            # Lender value chart
            st.plotly_chart(
                create_lender_value_chart(lenders),
                use_container_width=True
            )

            # All lenders table
            st.markdown("### All Lenders")
            df_lenders = pd.DataFrame(lenders)
            st.dataframe(
                df_lenders,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "estimated_value": st.column_config.NumberColumn(
                        "Estimated Value",
                        format="£%.2f"
                    ),
                    "avg_claim_value": st.column_config.NumberColumn(
                        "Avg Claim Value",
                        format="£%.2f"
                    ),
                    "pct_of_total": st.column_config.NumberColumn(
                        "% of Total",
                        format="%.2f%%"
                    )
                }
            )

        with tab6:
            st.markdown("## Raw Extracted Data")

            st.markdown("### Portfolio Metrics")
            st.json(portfolio)

            st.markdown("### Financial Metrics")
            st.json(monthly_data.get('financial_metrics', {}))

            st.markdown("### Pipeline Status")
            st.json(monthly_data.get('pipeline', {}))

            st.markdown("### Forecasting")
            st.json(monthly_data.get('forecasting', {}))

            st.markdown("### Profit Distribution Rules")
            st.json(profit_rules)

            st.markdown("### Full Investor Report Data")
            st.json(investor_report)

    else:
        # Show info when no file uploaded
        st.info("👆 Upload a monthly Excel report and click 'Generate Investor Report' to start the AI analysis")

        st.markdown("---")

        st.markdown("### 🤖 How It Works")

        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown("""
            **1️⃣ Priority Deed Agent**

            Reads the Priority Deed document and extracts:
            - Profit split percentages
            - DBA rate (30%)
            - Cost recovery rules
            - Payment waterfall
            """)

        with col2:
            st.markdown("""
            **2️⃣ FCA Compliance Agent**

            Reads FCA Redress Scheme and validates:
            - Commission thresholds
            - Eligible products
            - Required disclosures
            - Compliance criteria
            """)

        with col3:
            st.markdown("""
            **3️⃣ Monthly Report Agent**

            Analyzes Excel report to extract:
            - Portfolio metrics
            - All lender data (50-70 lenders)
            - Pipeline breakdown
            - Financial costs
            - Forecasting data
            """)

        with col4:
            st.markdown("""
            **4️⃣ Investor Report Agent**

            Combines all insights to generate:
            - Executive summary
            - Financial analysis
            - Compliance assessment
            - Risk analysis
            - Action items
            """)

    # Footer
    st.markdown("---")
    st.markdown("*Powered by OpenAI GPT-4o • 4-Agent Intelligent System • Factual Data-Driven Analysis*")
